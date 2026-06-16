from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "出货检验报告.docx"


def set_cell_text(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return paragraph


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="000000", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_cm):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_jc(table, value="center"):
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.first_child_found_in("w:jc")
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), value)


def set_column_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_centered_run(paragraph, text, size, bold=False, east_asia="宋体"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_kv_line(doc, items, size=8):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, 0, 1, 1.0)
    for i, (label, value, width_spaces) in enumerate(items):
        run = paragraph.add_run(label)
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(size)
        value_run.font.name = "Times New Roman"
        value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if i < len(items) - 1:
            paragraph.add_run(" " * width_spaces)


def add_signoff_block(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(245)
    spacer.paragraph_format.line_spacing = 1.0

    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_jc(table, "center")
    set_column_widths(table, [3.4, 3.5, 3.8, 5.6, 8.2])
    set_table_width(table, 24.5)

    labels = [
        "检验员\nOperated By",
        "审核\nApproved By",
        "日期\nDate",
        "报告编号\nRep No",
        "页码",
    ]
    values = [
        "梁起煜",
        "合 格",
        "2025-12-24",
        "PHW2025122411311",
        "第  1  页    共  1  页",
    ]

    for c, label in enumerate(labels):
        set_cell_text(table.cell(0, c), label, size=7.0)
        set_cell_text(table.cell(1, c), values[c], size=7.8, bold=(c == 1))
        for r in range(2):
            cell = table.cell(r, c)
            set_cell_margins(cell, top=20, bottom=20, start=20, end=20)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = borders.find(qn(f"w:{edge}"))
                if border is None:
                    border = OxmlElement(f"w:{edge}")
                    borders.append(border)
                border.set(qn("w:val"), "nil")
        if c == 1:
            run = table.cell(1, c).paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(190, 0, 0)
            run.font.size = Pt(10)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9)

    p = doc.add_paragraph()
    add_centered_run(p, "广东风华高新科技股份有限公司", 15.5, True, "黑体")
    p = doc.add_paragraph()
    add_centered_run(p, "FENGHUA  ADVANCED  TECHNOLOGY  HOLDING  CO., LTD.", 9.8, True)
    p = doc.add_paragraph()
    add_centered_run(p, "出货检验报告", 11.3, True, "黑体")
    p.add_run("  ")
    add_centered_run(p, "OUTGOING  INSPECTION  SHEET", 10.2, True)

    add_kv_line(
        doc,
        [
            ("客户名称 (Customer): ", "__________________________", 10),
            ("日期 (Date): ", "2025-12-24", 10),
            ("检验条件 (Condition): ", "IR: 22℃ / 50%", 10),
            ("可焊性条件 (Solder ability): ", "245±5℃  2±0.5S", 0),
        ],
        size=7.0,
    )

    widths = [2.2, 3.25, 2.65, 1.55, 1.75, 1.35, 1.5, 2.85, 1.35, 2.15, 1.25]
    table = doc.add_table(rows=4, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_jc(table, "center")
    set_table_width(table, sum(widths))
    set_table_borders(table)
    set_column_widths(table, widths)

    headers = [
        "客户型号\n(Customer P/N)",
        "型号规格\n(Part No.)",
        "批号\nLot No.",
        "绝缘电阻\nIR\nSpec.",
        "容量\nCap.\nRange",
        "损耗\nD.F.\nSpec.",
        "耐电压\nDWV\nSpec.",
        "容差/脚距\nCap. D/Y (C)",
        "额定电压\nR.V.",
        "耐压电流\nDWV",
        "结果\nResult",
    ]
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(0, idx), text, size=6.3, bold=True)
        set_cell_shading(table.cell(0, idx), "F2F2F2")
        set_cell_margins(table.cell(0, idx), top=45, bottom=45, start=55, end=55)

    rows = [
        [
            "DR00319",
            "0805B104K500NT",
            "NDX154983HMC",
            "IR ≥ 2GΩ",
            "90~110nF",
            "≤3.5%",
            "≥2.5Vr",
            "MMZ±10% / 1.0±0.2mm",
            "50V",
            "2.5Vr-50mA\nMAX",
            "合格",
        ],
        [
            "DR00319",
            "0805B104K500NT",
            "NDX171705N9D",
            "IR ≥ 2GΩ",
            "90~110nF",
            "≤3.7%",
            "≥2.5Vr",
            "MMZ±10% / 1.0±0.2mm",
            "50V",
            "2.5Vr-50mA\nMAX",
            "合格",
        ],
    ]
    for r_idx, data in enumerate(rows, start=1):
        for c_idx, text in enumerate(data):
            set_cell_text(table.cell(r_idx, c_idx), text, size=6.8)
            set_cell_margins(table.cell(r_idx, c_idx), top=58, bottom=58, start=55, end=55)

    for idx in range(len(widths)):
        set_cell_text(table.cell(3, idx), "", size=7)
        set_cell_margins(table.cell(3, idx), top=72, bottom=72, start=55, end=55)

    add_signoff_block(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
